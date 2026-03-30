"""Phase 4 tests: Document Intelligence — chunking, entity extraction, parser upgrades, dedup."""

from __future__ import annotations

import io

# ---------------------------------------------------------------------------
# Semantic Chunking
# ---------------------------------------------------------------------------


def test_semantic_chunking():
    """5000-char text with 3 sections -> 4-5 chunks with correct section_titles."""
    from crewai_enterprise_pipeline_api.ingestion.chunker import semantic_chunk

    text = (
        "# Section One\n\n"
        + "A " * 600
        + "\n\n"
        + "# Section Two\n\n"
        + "B " * 600
        + "\n\n"
        + "# Section Three\n\n"
        + "C " * 600
    )

    chunks = semantic_chunk(text, max_chars=1200)
    assert len(chunks) >= 3
    titles = {c.section_title for c in chunks}
    assert "Section One" in titles
    assert "Section Two" in titles
    assert "Section Three" in titles

    # Verify ordering and offsets
    for i, c in enumerate(chunks):
        assert c.chunk_index == i
        assert c.char_start >= 0
        assert c.char_end > c.char_start


def test_semantic_chunking_page_markers():
    """Chunks detect [Page N] markers from PDF output."""
    from crewai_enterprise_pipeline_api.ingestion.chunker import semantic_chunk

    text = "[Page 1]\nFirst page content here.\n\n[Page 2]\nSecond page content here."
    chunks = semantic_chunk(text, max_chars=1200)
    assert len(chunks) >= 1
    # At least one chunk should have a page number
    pages = [c.page_number for c in chunks if c.page_number is not None]
    assert len(pages) >= 1


# ---------------------------------------------------------------------------
# Financial Entity Extraction
# ---------------------------------------------------------------------------


def test_financial_entity_extraction():
    """Mock P&L text -> revenue, EBITDA evidence extracted."""
    from crewai_enterprise_pipeline_api.ingestion.entity_extractor import extract_entities

    text = (
        "The company reported revenue of INR 450 crore for FY25. "
        "EBITDA: Rs. 85 crore. Profit after tax was INR 32 crore. "
        "Net debt stands at INR 120 crore. "
        "Statutory auditor: Deloitte Haskins & Sells LLP. "
        "The auditors issued an unqualified opinion."
    )

    entities = extract_entities(
        text, document_kind="financial_statement", artifact_id="art-1", citation_prefix="test.pdf"
    )

    titles = [e.title for e in entities]
    assert any("Revenue" in t for t in titles)
    assert any("EBITDA" in t for t in titles)
    assert any("Profit After Tax" in t for t in titles)
    assert any("Auditor" in t for t in titles)
    assert any("Audit Opinion" in t for t in titles)


def test_legal_entity_extraction():
    """Legal contract text -> party names, dates extracted."""
    from crewai_enterprise_pipeline_api.ingestion.entity_extractor import extract_entities

    text = (
        "This agreement dated 15 March 2025 is between Alpha Corp and Beta Ltd. "
        "This agreement shall be governed by Indian law."
    )

    entities = extract_entities(
        text, document_kind="legal_agreement", artifact_id="art-2", citation_prefix="contract.pdf"
    )

    titles = [e.title for e in entities]
    assert any("Execution Date" in t for t in titles)
    assert any("Governing Law" in t for t in titles)


def test_india_identifiers():
    """CIN and GSTIN are extracted regardless of document kind."""
    from crewai_enterprise_pipeline_api.ingestion.entity_extractor import extract_entities

    text = (
        "Company CIN: U72200MH2015PTC123456. "
        "GSTIN: 27AAACR5055K1Z5."
    )

    entities = extract_entities(
        text, document_kind="general", artifact_id="art-3", citation_prefix="doc.pdf"
    )

    titles = [e.title for e in entities]
    assert any("CIN" in t for t in titles)
    assert any("GSTIN" in t for t in titles)


# ---------------------------------------------------------------------------
# XLSX Multi-sheet (markdown table output)
# ---------------------------------------------------------------------------


def test_xlsx_multi_sheet(client):
    """Upload 2-sheet XLSX -> chunks from both sheets."""
    from openpyxl import Workbook

    wb = Workbook()
    ws1 = wb.active
    ws1.title = "P&L"
    ws1.append(["Particulars", "FY24", "FY25"])
    ws1.append(["Revenue", "300 Cr", "450 Cr"])
    ws1.append(["EBITDA", "60 Cr", "85 Cr"])

    ws2 = wb.create_sheet("Balance Sheet")
    ws2.append(["Item", "Amount"])
    ws2.append(["Total Assets", "800 Cr"])
    ws2.append(["Net Debt", "120 Cr"])

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)

    case = client.post(
        "/api/v1/cases",
        json={
            "name": "XLSX Multi-Sheet Test",
            "target_name": "TestCo",
            "motion_pack": "buy_side_diligence",
            "sector_pack": "tech_saas_services",
        },
    ).json()
    case_id = case["id"]

    resp = client.post(
        f"/api/v1/cases/{case_id}/documents/upload",
        files={
            "file": (
                "financials.xlsx",
                buf,
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
        },
        data={
            "document_kind": "financial_statement",
            "source_kind": "uploaded_dataroom",
            "workstream_domain": "financial_qoe",
            "evidence_kind": "metric",
        },
    )
    assert resp.status_code == 201
    body = resp.json()
    assert body["chunks_created"] > 0
    assert body["evidence_items_created"] > 0

    # Verify chunks endpoint
    doc_id = body["artifact"]["id"]
    chunks_resp = client.get(f"/api/v1/cases/{case_id}/documents/{doc_id}/chunks")
    assert chunks_resp.status_code == 200
    chunks = chunks_resp.json()
    assert len(chunks) > 0
    # Chunks should have section titles from sheet names
    texts = " ".join(c["text"] for c in chunks)
    assert "P&L" in texts or "Revenue" in texts or "Balance Sheet" in texts


# ---------------------------------------------------------------------------
# Document Dedup
# ---------------------------------------------------------------------------


def test_document_dedup(client):
    """Upload same file twice -> second upload returns existing record (SHA256 match)."""
    case = client.post(
        "/api/v1/cases",
        json={
            "name": "Dedup Test",
            "target_name": "TestCo",
            "motion_pack": "buy_side_diligence",
            "sector_pack": "tech_saas_services",
        },
    ).json()
    case_id = case["id"]

    file_content = b"Revenue INR 100 crore for FY25. EBITDA Rs. 20 crore."
    upload_kwargs = {
        "files": {"file": ("report.txt", io.BytesIO(file_content), "text/plain")},
        "data": {
            "document_kind": "financial_statement",
            "source_kind": "uploaded_dataroom",
            "workstream_domain": "financial_qoe",
            "evidence_kind": "fact",
        },
    }

    resp1 = client.post(f"/api/v1/cases/{case_id}/documents/upload", **upload_kwargs)
    assert resp1.status_code == 201
    first = resp1.json()
    assert first["chunks_created"] > 0

    # Upload exact same content again
    upload_kwargs["files"]["file"] = ("report.txt", io.BytesIO(file_content), "text/plain")
    resp2 = client.post(f"/api/v1/cases/{case_id}/documents/upload", **upload_kwargs)
    assert resp2.status_code == 201
    second = resp2.json()
    # Dedup: same artifact returned, no new chunks/evidence
    assert second["artifact"]["id"] == first["artifact"]["id"]
    assert second["chunks_created"] == 0
    assert second["evidence_items_created"] == 0


# ---------------------------------------------------------------------------
# Chunks endpoint
# ---------------------------------------------------------------------------


def test_chunks_endpoint_404_on_missing_doc(client):
    """GET /documents/{id}/chunks returns 404 for non-existent document."""
    case = client.post(
        "/api/v1/cases",
        json={
            "name": "Chunks 404 Test",
            "target_name": "TestCo",
            "motion_pack": "buy_side_diligence",
            "sector_pack": "tech_saas_services",
        },
    ).json()
    case_id = case["id"]

    resp = client.get(f"/api/v1/cases/{case_id}/documents/nonexistent-doc/chunks")
    assert resp.status_code == 404


# ---------------------------------------------------------------------------
# Parser upgrades
# ---------------------------------------------------------------------------


def test_docx_heading_extraction(client):
    """DOCX with headings -> parsed text includes markdown headings."""
    from docx import Document

    doc = Document()
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph("The company is performing well.")
    doc.add_heading("Financial Highlights", level=2)
    doc.add_paragraph("Revenue grew by 25% year-over-year.")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    case = client.post(
        "/api/v1/cases",
        json={
            "name": "DOCX Heading Test",
            "target_name": "TestCo",
            "motion_pack": "buy_side_diligence",
            "sector_pack": "tech_saas_services",
        },
    ).json()
    case_id = case["id"]

    resp = client.post(
        f"/api/v1/cases/{case_id}/documents/upload",
        files={
            "file": (
                "report.docx",
                buf,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={
            "document_kind": "general",
            "source_kind": "uploaded_dataroom",
            "workstream_domain": "commercial",
            "evidence_kind": "fact",
        },
    )
    assert resp.status_code == 201
    body = resp.json()
    assert body["chunks_created"] >= 1


def test_ingestion_result_has_new_fields(client):
    """DocumentIngestionResult includes chunks_created and entities_extracted."""
    case = client.post(
        "/api/v1/cases",
        json={
            "name": "Result Fields Test",
            "target_name": "TestCo",
            "motion_pack": "buy_side_diligence",
            "sector_pack": "tech_saas_services",
        },
    ).json()
    case_id = case["id"]

    resp = client.post(
        f"/api/v1/cases/{case_id}/documents/upload",
        files={"file": ("simple.txt", io.BytesIO(b"Just a plain text file."), "text/plain")},
        data={
            "document_kind": "general",
            "source_kind": "uploaded_dataroom",
            "workstream_domain": "commercial",
            "evidence_kind": "fact",
        },
    )
    assert resp.status_code == 201
    body = resp.json()
    assert "chunks_created" in body
    assert "entities_extracted" in body

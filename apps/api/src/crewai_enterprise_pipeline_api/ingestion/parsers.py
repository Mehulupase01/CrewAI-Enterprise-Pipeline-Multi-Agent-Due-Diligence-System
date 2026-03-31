from __future__ import annotations

import csv
import io
import json
import logging
from dataclasses import dataclass

import pdfplumber
from docx import Document as DocxDocument
from openpyxl import load_workbook

logger = logging.getLogger(__name__)


@dataclass
class ParsedDocument:
    parser_name: str
    text: str

    @property
    def extracted_character_count(self) -> int:
        return len(self.text)


class DocumentParser:
    def parse(self, filename: str, mime_type: str | None, content: bytes) -> ParsedDocument:
        lower_name = filename.lower()
        if mime_type in {"text/plain", "text/markdown"} or lower_name.endswith((".txt", ".md")):
            return ParsedDocument(parser_name="plaintext", text=self._decode_text(content))
        if mime_type == "application/json" or lower_name.endswith(".json"):
            return ParsedDocument(parser_name="json", text=self._parse_json(content))
        if mime_type in {"text/csv", "application/csv"} or lower_name.endswith(".csv"):
            return ParsedDocument(parser_name="csv", text=self._parse_csv(content))
        if mime_type == "application/pdf" or lower_name.endswith(".pdf"):
            return ParsedDocument(parser_name="pdfplumber", text=self._parse_pdf(content))
        if lower_name.endswith((".xlsx", ".xlsm")):
            return ParsedDocument(parser_name="openpyxl", text=self._parse_xlsx(content))
        if mime_type in {
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        } or lower_name.endswith(".docx"):
            return ParsedDocument(parser_name="python-docx", text=self._parse_docx(content))
        return ParsedDocument(parser_name="binary-fallback", text=self._decode_text(content))

    def _decode_text(self, content: bytes) -> str:
        for encoding in ("utf-8", "utf-16", "latin-1"):
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        return content.decode("utf-8", errors="ignore")

    def _parse_json(self, content: bytes) -> str:
        try:
            loaded = json.loads(self._decode_text(content))
            return json.dumps(loaded, indent=2, ensure_ascii=False)
        except Exception:
            logger.warning("JSON parse failed", exc_info=True)
            return ""

    def _parse_csv(self, content: bytes) -> str:
        try:
            buffer = io.StringIO(self._decode_text(content))
            reader = csv.reader(buffer)
            rows = [" | ".join(cell.strip() for cell in row) for row in reader]
            return "\n".join(rows)
        except Exception:
            logger.warning("CSV parse failed", exc_info=True)
            return ""

    def _parse_pdf(self, content: bytes) -> str:
        try:
            buffer = io.BytesIO(content)
            pages: list[str] = []
            with pdfplumber.open(buffer) as pdf:
                for page_number, page in enumerate(pdf.pages, start=1):
                    parts: list[str] = []
                    page_text = page.extract_text() or ""
                    if page_text.strip():
                        parts.append(page_text.strip())

                    # Extract tables as markdown
                    try:
                        tables = page.extract_tables()
                        for table in tables or []:
                            md = _table_to_markdown(table)
                            if md:
                                parts.append(md)
                    except Exception:
                        pass  # table extraction is best-effort

                    if parts:
                        pages.append(f"[Page {page_number}]\n" + "\n\n".join(parts))
            return "\n\n".join(pages)
        except Exception:
            logger.warning("PDF parse failed", exc_info=True)
            return ""

    def _parse_xlsx(self, content: bytes) -> str:
        try:
            buffer = io.BytesIO(content)
            workbook = load_workbook(buffer, data_only=True)
            sheets: list[str] = []
            for worksheet in workbook.worksheets[:8]:
                rows: list[list[str]] = []
                for row in worksheet.iter_rows(values_only=True, max_row=200):
                    values = [str(v).strip() if v is not None else "" for v in row]
                    if any(v for v in values):
                        rows.append(values)
                if not rows:
                    continue
                md = _rows_to_markdown_table(rows)
                sheets.append(f"## {worksheet.title}\n\n{md}")
            return "\n\n".join(sheets)
        except Exception:
            logger.warning("XLSX parse failed", exc_info=True)
            return ""

    def _parse_docx(self, content: bytes) -> str:
        try:
            buffer = io.BytesIO(content)
            document = DocxDocument(buffer)
            parts: list[str] = []
            for paragraph in document.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    continue
                style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
                if "heading" in style_name:
                    level = 1
                    for ch in style_name:
                        if ch.isdigit():
                            level = int(ch)
                            break
                    parts.append(f"{'#' * level} {text}")
                else:
                    parts.append(text)

            # Extract tables from DOCX
            for table in document.tables:
                rows: list[list[str]] = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(c for c in cells):
                        rows.append(cells)
                if rows:
                    parts.append(_rows_to_markdown_table(rows))

            return "\n\n".join(parts)
        except Exception:
            logger.warning("DOCX parse failed", exc_info=True)
            return ""


def _table_to_markdown(table: list[list[str | None]]) -> str:
    """Convert a pdfplumber table (list of rows) to markdown table format."""
    if not table:
        return ""
    rows: list[list[str]] = []
    for row in table:
        cells = [str(c).strip() if c else "" for c in row]
        if any(c for c in cells):
            rows.append(cells)
    return _rows_to_markdown_table(rows)


def _rows_to_markdown_table(rows: list[list[str]]) -> str:
    """Convert a list of string rows to a markdown table."""
    if not rows:
        return ""
    # Use first row as header
    header = rows[0]
    col_count = len(header)
    lines = ["| " + " | ".join(header) + " |"]
    lines.append("| " + " | ".join("---" for _ in range(col_count)) + " |")
    for row in rows[1:]:
        # Pad or truncate to match header column count
        padded = row[:col_count] + [""] * max(0, col_count - len(row))
        lines.append("| " + " | ".join(padded) + " |")
    return "\n".join(lines)


def chunk_text(text: str, *, max_chars: int = 1200) -> list[str]:
    cleaned = "\n".join(line.rstrip() for line in text.splitlines()).strip()
    if not cleaned:
        return []

    paragraphs = [paragraph.strip() for paragraph in cleaned.split("\n\n") if paragraph.strip()]
    chunks: list[str] = []
    current = ""

    for paragraph in paragraphs:
        candidate = f"{current}\n\n{paragraph}".strip() if current else paragraph
        if len(candidate) <= max_chars:
            current = candidate
            continue
        if current:
            chunks.append(current)
        if len(paragraph) <= max_chars:
            current = paragraph
            continue
        start = 0
        while start < len(paragraph):
            end = start + max_chars
            chunks.append(paragraph[start:end])
            start = end
        current = ""

    if current:
        chunks.append(current)

    return chunks

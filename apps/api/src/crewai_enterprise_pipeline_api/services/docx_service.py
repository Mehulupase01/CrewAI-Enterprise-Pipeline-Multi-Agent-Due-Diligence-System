from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from crewai_enterprise_pipeline_api.services.report_markdown import (
    BulletListBlock,
    HeadingBlock,
    ParagraphBlock,
    TableBlock,
    collect_toc_headings,
    parse_markdown_blocks,
)


class DocxService:
    def render_report(
        self,
        *,
        title: str,
        subtitle: str,
        template_label: str,
        markdown: str,
    ) -> bytes:
        document = Document()
        self._configure_styles(document)

        self._add_cover_page(
            document,
            title=title,
            subtitle=subtitle,
            template_label=template_label,
        )

        blocks = parse_markdown_blocks(markdown)
        toc_headings = [heading for heading in collect_toc_headings(blocks) if heading != title]
        if toc_headings:
            document.add_heading("Table of Contents", level=1)
            for heading in toc_headings:
                document.add_paragraph(heading, style="List Bullet")
            document.add_page_break()

        for block in blocks:
            if isinstance(block, HeadingBlock):
                document.add_heading(block.text, level=min(block.level, 4))
            elif isinstance(block, ParagraphBlock):
                document.add_paragraph(block.text)
            elif isinstance(block, BulletListBlock):
                for item in block.items:
                    document.add_paragraph(item, style="List Bullet")
            elif isinstance(block, TableBlock):
                self._add_table(document, block)

        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    def _configure_styles(self, document: Document) -> None:
        normal_style = document.styles["Normal"]
        normal_style.font.name = "Calibri"
        normal_style.font.size = Pt(10.5)

    def _add_cover_page(
        self,
        document: Document,
        *,
        title: str,
        subtitle: str,
        template_label: str,
    ) -> None:
        title_paragraph = document.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(22)

        label_paragraph = document.add_paragraph()
        label_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label_run = label_paragraph.add_run(template_label)
        label_run.italic = True
        label_run.font.size = Pt(12)

        subtitle_paragraph = document.add_paragraph()
        subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle_paragraph.add_run(subtitle)
        subtitle_run.font.size = Pt(10)

        document.add_page_break()

    def _add_table(self, document: Document, block: TableBlock) -> None:
        if not block.headers:
            return

        table = document.add_table(rows=1, cols=len(block.headers))
        table.style = "Table Grid"
        header_cells = table.rows[0].cells
        for idx, header in enumerate(block.headers):
            header_cells[idx].text = header

        for row in block.rows:
            row_cells = table.add_row().cells
            for idx, value in enumerate(row):
                row_cells[idx].text = value
